import os
import docx

def update_proposal(file_path):
    print(f"Updating {file_path}")
    doc = docx.Document(file_path)
    
    # 1. Update Section 1.1.2 State of the Art
    for para in doc.paragraphs:
        if "TRL progression:" in para.text:
            para.text = "TRL progression: entry TRL 5 (concept validated and MVP fully integrated with COP-PILOT NGSI-LD and TMForum standards) → exit TRL 8. Novel contribution to COP-PILOT: first integration of a public Layer-1 blockchain (Solana) as an immutable data anchoring layer for IoT provenance records within a COP-PILOT service. The service is already containerized (Docker), orchestrated (Helm), and exposes an LLM-assisted service onboarding via TMF 633 (Service Catalog) and SIF OpenZiti mTLS micro-tunneling."
            break
            
    # 2. Update Section 1.1.4 Data Management if it's there
    for i, para in enumerate(doc.paragraphs):
        if "1.1.4  Data Management" in para.text:
            if i + 1 < len(doc.paragraphs) and doc.paragraphs[i+1].text.strip() == "":
                doc.paragraphs[i+1].text = "AgroChain has fully implemented the ETSI NGSI-LD 'GrainLot' data model, adhering to FIWARE Smart Data Models. The FastAPI backend exposes an API endpoint allowing the Orion-LD Context Broker to immediately ingest cross-border telemetry (GPS, Cadastre, UKAS CertCheck) and LLM portal agents to run semantic search over Ukrainian agro-parcels."
            break
            
    # 3. Update Technical Impact
    for para in doc.paragraphs:
        if "AgroChain Ukraine validates four COP-PILOT innovations" in para.text:
            para.text = "AgroChain Ukraine has successfully validated and implemented the core COP-PILOT structural requirements: (1) LLM-assisted service onboarding — our TMF 633 JSON payload configures OpenSlice for native natural language querying; (2) SIF zero-trust cross-border tunnels — Helm charts natively deploy OpenZiti identities; (3) NGSI-LD multi-domain data sharing — our custom translator module translates Solana transaction hashes and Derzhprodspozhyvsluzhba Phytosanitary statuses into strict ETSI NGSI-LD standards. (4) CI/CD GitOps workflow is structurally built for push-button platform integration."
            break
            
    doc.save(file_path.replace(".docx", "_updated.docx"))
    print("Proposal updated.")

def update_ethics(file_path):
    print(f"Updating {file_path}")
    doc = docx.Document(file_path)
    # Check if we can add a note about PII hashing
    for para in doc.paragraphs:
        if "Personal Data" in para.text or "processing of personal data" in para.text.lower():
            para.text += " Note: AgroChain Ukraine strictly hashes all Personally Identifiable Information (PII) such as the Farmer's Tax ID (RNOKPP) using cryptographic SHA-256 hashing before any record hits the NGSI-LD Context Broker or public Blockchain, ensuring compliance with EU GDPR and Annex III Ethics framework."
            
    doc.save(file_path.replace(".docx", "_updated.docx"))
    print("Ethics updated.")

proposal_path = "E:/Hackatones/world-bank/files/AGROCHAIN-UA_CPOC1_Proposal_20260504.docx"
ethics_path = "E:/Hackatones/world-bank/files/AGROCHAIN-UA_CPOC1_Annexes_DoH_SME_Ethics.docx"

if os.path.exists(proposal_path):
    update_proposal(proposal_path)
if os.path.exists(ethics_path):
    update_ethics(ethics_path)
